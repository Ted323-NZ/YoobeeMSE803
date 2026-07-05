"""DOCX rendering for optimized CV text."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path


SECTION_HEADINGS = {
    "PROFESSIONAL SUMMARY",
    "TECHNICAL SKILLS",
    "PROJECTS",
    "EXPERIENCE",
    "ACADEMIC EXPERIENCE",
    "EDUCATION",
    "CERTIFICATIONS AND DEVELOPMENT",
    "CERTIFICATIONS",
}


def build_cv_docx_bytes(optimized_cv_text: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.25

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25

    lines = [line.rstrip() for line in optimized_cv_text.strip().splitlines()]
    header_lines: list[str] = []
    body_lines: list[str] = []
    in_body = False
    for line in lines:
        if line.strip().upper() in SECTION_HEADINGS:
            in_body = True
        if in_body:
            body_lines.append(line)
        else:
            header_lines.append(line)

    if header_lines:
        name = header_lines[0].strip()
        role = header_lines[1].strip() if len(header_lines) > 1 else ""
        contact = header_lines[2].strip() if len(header_lines) > 2 else ""
        name_paragraph = document.add_paragraph()
        name_paragraph.paragraph_format.space_after = Pt(2)
        run = name_paragraph.add_run(name)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

        if role:
            role_paragraph = document.add_paragraph()
            role_paragraph.paragraph_format.space_after = Pt(2)
            role_run = role_paragraph.add_run(role)
            role_run.bold = True
            role_run.font.name = "Calibri"
            role_run.font.size = Pt(12)
            role_run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)

        if contact:
            contact_paragraph = document.add_paragraph()
            contact_paragraph.paragraph_format.space_after = Pt(8)
            contact_run = contact_paragraph.add_run(contact)
            contact_run.font.name = "Calibri"
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for raw_line in body_lines:
        line = raw_line.strip()
        if not line:
            continue
        if line.upper() in SECTION_HEADINGS:
            document.add_heading(line.title(), level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def write_cv_docx(optimized_cv_text: str, output_path: str | Path) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(build_cv_docx_bytes(optimized_cv_text))
    return path
